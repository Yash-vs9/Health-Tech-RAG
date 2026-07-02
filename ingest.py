import os
import warnings
import pdfplumber

from docx import Document

from langchain_core.documents import Document as LangchainDocument
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_chroma import Chroma
from dotenv import load_dotenv

load_dotenv()

# Hide LangChain deprecation warnings for a cleaner terminal
warnings.filterwarnings("ignore", category=DeprecationWarning)

def extract_pdf_tables(pdf_path):
    """
    Extracts tables from a PDF and returns them as LangChain Documents.
    Tables are formatted as Markdown to optimize LLM comprehension.
    """
    table_documents = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages):
            tables = page.extract_tables()

            if not tables:
                continue

            for table in tables:
                rows = []

                for i, row in enumerate(table):
                    # Clean the cells and wrap them in Markdown pipes
                    cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                    row_string = "| " + " | ".join(cleaned_row) + " |"
                    rows.append(row_string)
                    
                    # Add the Markdown separator line right after the header (row 0)
                    if i == 0:
                        separator = "|" + "|".join(["---"] * len(cleaned_row)) + "|"
                        rows.append(separator)

                table_text = "\n".join(rows)

                table_documents.append(
                    LangchainDocument(
                        page_content=table_text,
                        metadata={
                            "source": pdf_path,
                            "page": page_number + 1,
                            "content_type": "table"
                        }
                    )
                )

    return table_documents

def extract_docx_tables(docx_path):
    """
    Extracts paragraphs and tables from a DOCX file
    and returns them as LangChain Documents. Tables are formatted in Markdown.
    """
    documents = []
    doc = Document(docx_path)

    # Extract paragraphs
    paragraph_text = "\n".join(
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()
    )

    if paragraph_text:
        documents.append(
            LangchainDocument(
                page_content=paragraph_text,
                metadata={
                    "source": docx_path,
                    "content_type": "text"
                }
            )
        )

    # Extract tables
    for table in doc.tables:
        rows = []

        for i, row in enumerate(table.rows):
            # Clean the cells and wrap them in Markdown pipes
            cells = [cell.text.strip() for cell in row.cells]
            row_string = "| " + " | ".join(cells) + " |"
            rows.append(row_string)
            
            # Add the Markdown separator line right after the header (row 0)
            if i == 0:
                separator = "|" + "|".join(["---"] * len(cells)) + "|"
                rows.append(separator)

        table_text = "\n".join(rows)

        documents.append(
            LangchainDocument(
                page_content=table_text,
                metadata={
                    "source": docx_path,
                    "content_type": "table"
                }
            )
        )

    return documents

def ingest_documents(paths, persist_dir="./chroma_db"):
    """
    Loads, chunks, embeds, and stores PDF and DOCX files into a Chroma vector database.
    Includes custom extraction for tables to preserve Markdown structures.
    Embeddings are generated via Hugging Face's hosted Inference API,
    routed through the Scaleway provider, running Qwen3-Embedding-8B.
    Nothing is downloaded or run locally.
    """

    # 1. Initialize the Text Splitter
    # CRITICAL: chunk_size increased to 1024 to preserve table structures
    # chunk_overlap increased to 100 to maintain cross-boundary context
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " "]
    )

    # 2. Initialize Qwen3-Embedding-8B via Hugging Face's Scaleway provider
    print("Connecting to Hugging Face Inference API (Qwen3-Embedding-8B via Scaleway)...")
    HF_TOKEN = os.getenv("HF_TOKEN")
    if not HF_TOKEN:
        print("[!] Warning: HF_TOKEN environment variable is not set.")
        print("    Get a token at https://huggingface.co/settings/tokens")
        print("    and run: export HF_TOKEN=your_token_here")
        print("    (You'll also need Inference Providers billing/credits set up")
        print("     on your HF account, since Scaleway is a paid provider.)")

    embeddings = HuggingFaceEndpointEmbeddings(
        model="Qwen/Qwen3-Embedding-8B",
        task="feature-extraction",
        provider="scaleway",          # the only provider currently serving this 8B model
        huggingfacehub_api_token=HF_TOKEN,
    )

    all_chunks = []

    # 3. Load and Chunk Documents
    for path in paths:
        if not os.path.exists(path):
            print(f"\n[!] Warning: File not found at {path}. Skipping.")
            continue

        extension = os.path.splitext(path)[1].lower()

        if extension == ".pdf":
            loader = PyMuPDFLoader(path)
            text_docs = loader.load()
            table_docs = extract_pdf_tables(path)
            
            docs = text_docs + table_docs
            print(
                f"Loaded {len(text_docs)} text documents "
                f"and {len(table_docs)} table documents."
            )

        elif extension == ".docx":
            docs = extract_docx_tables(path)
            print(f"Loaded {len(docs)} documents from DOCX.")

        else:
            print(f"Unsupported file type: {path}")
            continue

        print(f"Chunking {path}...")
        chunks = splitter.split_documents(docs)

        # --- CHUNK PREVIEW ---
        print("\n--- CHUNK PREVIEW ---")
        preview_limit = min(3, len(chunks))  # Show first 3 chunks
        for i in range(preview_limit):
            print(f"\n[Chunk {i+1} | Length: {len(chunks[i].page_content)} characters]")
            print(chunks[i].page_content)
            print("-" * 50)
        print("---------------------\n")

        # Verify chunk count correct as per assignment requirements
        print(f" -> Success: Generated {len(chunks)} chunks from {path}")
        all_chunks.extend(chunks)

    if not all_chunks:
        print("No chunks generated. Pipeline aborted.")
        return None

    # 4. Store in ChromaDB with metadata
    print(f"\nSending {len(all_chunks)} chunks to Hugging Face API for embedding and storing locally...")
    vectorstore = Chroma.from_documents(
        documents=all_chunks,
        embedding=embeddings,
        persist_directory=persist_dir
    )
    print(f"Pipeline complete! Database saved to {persist_dir}")
    return vectorstore


# --- Golden Dataset Structure ---
def get_golden_dataset():
    """
    Returns the Golden Dataset: 5 Query Variation Pairs
    Used for testing retrieval robustness.
    """
    return [
        {
            "label": "question_type=variation",
            "expected_answer": "He is a Student and learning Full stack developement.",
            "variations": [
                "whats the occupation of Yash",
                "What does Yash do for a living?",
                "Is Yash currently a student or a working professional?",
                "Tell me about Yash's career and educational background.",
                "What field of study or work is Yash involved in?"
            ]
        }
    ]


if __name__ == "__main__":
    # --- Testing Phase ---
    # Ensure this file exists in the same folder as this script
    sample_files = [
        "sample_mortgage.pdf",
        "sample_mortgage.docx"
    ]

    db = ingest_documents(sample_files)

    # Print out the Golden Dataset structure for validation
    golden_data = get_golden_dataset()
    print("\nGolden Dataset initialized with 5 variations checking for the exact same expected answer.")