"""
Document ingestion pipeline — PDF/DOCX/Image to Qdrant.

This is the core ingestion engine. It takes raw file bytes, extracts text
via multiple methods (parser, vision LLM, OCR), chunks the content, embeds
it, and stores it in Qdrant.

Pipeline:
    1. Load document (PyMuPDF for PDF, python-docx for DOCX, PIL for images)
    2. Extract tables as separate markdown chunks
    3. Vision LLM augmentation (PDF pages + DOCX embedded images)
    4. OCR fallback for scanned pages
    5. Merge parser + vision + OCR text
    6. Split into chunks (1024 chars, 50 overlap)
    7. Embed via Qwen3-Embedding-8B (4096-dim)
    8. Upsert to Qdrant in batches
    9. Refresh BM25 index

Functions:
    ingest_document(file_bytes, filename, doc_id) -> dict

    _load_pdf(file_path) -> list[Document]
    _load_docx(file_path) -> list[Document]
    _describe_chart_with_vision(image, filename) -> str | None
    _extract_vision_text_from_pdf(file_path) -> dict[int, str]
    _extract_vision_text_from_docx(file_path) -> list[str]
    _extract_tables_from_pdf(file_path) -> dict[int, list[str]]
    _extract_tables_from_docx(file_path) -> list[str]

Env vars used:
    CHUNK_SIZE, CHUNK_OVERLAP
    VISION_MODEL, VISION_TIMEOUT, VISION_DELAY
    VISION_PDF_ENABLED, VISION_PDF_PAGE_LIMIT, VISION_PDF_ZOOM
    VISION_DOCX_ENABLED, VISION_DOCX_IMAGE_LIMIT
"""

import os
import io
import base64
import hashlib
import uuid
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from backend.logging_config import get_logger
from .embeddings import get_embeddings
from . import vectorstore
from .ocr_fallback import apply_ocr_fallback
from PIL import Image
import pytesseract
from backend.services.retriever import refresh_bm25
from .upload_utils import safe_filename

logger = get_logger("backend.ingestion")


# Vision extraction via Vision LLM


def _get_vision_api_key() -> str | None:
    """Get NVIDIA API key for vision model."""
    nvidia_keys = os.getenv("NVIDIA_API_KEYS", os.getenv("NVIDIA_API_KEY", ""))
    if nvidia_keys:
        key = nvidia_keys.split(",")[0].strip()
        return key or None
    return None


def _image_to_base64_png(image: Image.Image) -> str:
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def _call_vision_model(image: Image.Image, prompt: str, context_name: str) -> str | None:
    """Call NVIDIA NIM vision model with an image + prompt."""
    api_key = _get_vision_api_key()
    if not api_key:
        logger.debug("No NVIDIA API key found for vision extraction (%s)", context_name)
        return None

    try:
        from langchain_core.messages import HumanMessage
        from langchain_nvidia_ai_endpoints import ChatNVIDIA

        img_base64 = _image_to_base64_png(image)
        message = HumanMessage(
            content=[
                {"type": "text", "text": prompt},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_base64}"},
                },
            ]
        )

        vision_model = os.getenv("VISION_MODEL", "nvidia/llama-3.1-nemotron-nano-vl-8b-v1")
        timeout = int(os.getenv("VISION_TIMEOUT", "120"))
        llm = ChatNVIDIA(
            model=vision_model,
            api_key=api_key,
            temperature=0.0,
            max_tokens=2048,
            timeout=timeout,
        )

        result = llm.invoke([message])
        content = (result.content or "").strip()
        if content:
            logger.debug("Vision extraction ok (%s) - chars=%d", context_name, len(content))
        return content or None
    except Exception as e:
        logger.warning("Vision model failed for %s: %s", context_name, e)
        return None


def _merge_extracted_and_vision_text(extracted_text: str, vision_text: str) -> str:
    """Merge parser/OCR text with vision output while avoiding exact duplication."""
    extracted = (extracted_text or "").strip()
    vision = (vision_text or "").strip()

    if not extracted:
        return vision
    if not vision:
        return extracted

    norm_extracted = " ".join(extracted.lower().split())
    norm_vision = " ".join(vision.lower().split())
    if norm_vision and norm_vision in norm_extracted:
        return extracted

    return (
        f"{extracted}\n\n"
        "[Vision Augmentation]\n"
        f"{vision}\n"
        "[/Vision Augmentation]"
    )


def _extract_document_text_with_vision(image: Image.Image, context_name: str) -> str | None:
    """Extract rich text from document/chart screenshots via vision LLM."""
    prompt = (
        "You are extracting content from a document image for retrieval. "
        "Return plain text with these sections if present: TITLE, HEADINGS, BODY TEXT, "
        "TABLES, KEY FIGURES, and SUMMARY. Include exact numbers and important entities. "
        "If the image has a chart, include chart type, axis labels, and each data point."
    )
    return _call_vision_model(image=image, prompt=prompt, context_name=context_name)


# ── Chart understanding via Vision LLM ─────────────────────────────────────


def _describe_chart_with_vision(image: Image.Image, filename: str) -> str | None:
    """
    Use NVIDIA NIM vision model to describe chart data from an image.
    Returns extracted data description or None if unavailable.
    """
    prompt = """Analyze this chart image and extract ALL numerical data. Format your response as:

CHART TYPE: [bar chart/line graph/pie chart/etc.]

TITLE: [chart title if visible]

AXES:
- X-axis: [label and values/categories]
- Y-axis: [label and values/range]

DATA:
[List each data point with exact values. For bar charts: category = value. For line charts: (x, y) pairs.]

SUMMARY: [1-2 sentence summary of what the chart shows]"""
    return _call_vision_model(image=image, prompt=prompt, context_name=f"chart:{filename}")


def _extract_vision_text_from_pdf(file_path: str) -> dict[int, str]:
    """
    Render PDF pages to images and extract supplemental text with the vision model.

    Uses parallel extraction with ThreadPoolExecutor for faster processing.
    Returns: {page_index: vision_text}
    """
    if os.getenv("VISION_PDF_ENABLED", "true").lower() != "true":
        return {}
    if not _get_vision_api_key():
        return {}

    import fitz

    page_limit = int(os.getenv("VISION_PDF_PAGE_LIMIT", "0"))
    zoom = float(os.getenv("VISION_PDF_ZOOM", "1.5"))
    max_workers = int(os.getenv("VISION_MAX_WORKERS", "3"))
    basename = os.path.basename(file_path)

    vision_by_page: dict[int, str] = {}
    pdf = fitz.open(file_path)
    try:
        total_pages = len(pdf)
        pages_to_process = total_pages if page_limit <= 0 else min(total_pages, page_limit)
        logger.info(
            "Vision PDF extraction - file=%s, pages=%d/%d, workers=%d",
            basename, pages_to_process, total_pages, max_workers,
        )

        def _extract_page(page_idx: int) -> tuple[int, str | None]:
            try:
                page = pdf[page_idx]
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                vision_text = _extract_document_text_with_vision(
                    image=image,
                    context_name=f"pdf:{basename}:page:{page_idx + 1}",
                )
                if vision_text and len(vision_text.strip()) >= 20:
                    return (page_idx, vision_text.strip())
            except Exception as e:
                logger.debug("Vision PDF extraction failed on page %d: %s", page_idx + 1, e)
            return (page_idx, None)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(_extract_page, idx): idx for idx in range(pages_to_process)}
            for future in as_completed(futures):
                page_idx, text = future.result()
                if text:
                    vision_by_page[page_idx] = text
    finally:
        pdf.close()

    logger.info(
        "Vision PDF extraction done - file=%s, pages_augmented=%d",
        basename, len(vision_by_page),
    )
    return vision_by_page


def _extract_vision_text_from_docx(file_path: str) -> list[str]:
    """
    Extract supplemental text from embedded DOCX images using vision model.

    Returns: list of extracted text blocks (one per image).
    """
    if os.getenv("VISION_DOCX_ENABLED", "true").lower() != "true":
        return []
    if not _get_vision_api_key():
        return []

    import docx

    image_limit = int(os.getenv("VISION_DOCX_IMAGE_LIMIT", "10"))
    basename = os.path.basename(file_path)

    vision_blocks: list[str] = []
    seen_hashes: set[str] = set()

    try:
        doc = docx.Document(file_path)
        image_blobs: list[bytes] = []
        for rel in doc.part.rels.values():
            target_ref = str(getattr(rel, "target_ref", ""))
            if "image" not in target_ref:
                continue
            blob = getattr(rel.target_part, "blob", None)
            if blob:
                image_blobs.append(blob)

        if image_limit > 0:
            image_blobs = image_blobs[:image_limit]

        for idx, blob in enumerate(image_blobs):
            digest = hashlib.sha256(blob).hexdigest()
            if digest in seen_hashes:
                continue
            seen_hashes.add(digest)

            try:
                image = Image.open(io.BytesIO(blob))
                vision_text = _extract_document_text_with_vision(
                    image=image,
                    context_name=f"docx:{basename}:image:{idx + 1}",
                )
                if vision_text and len(vision_text.strip()) >= 20:
                    vision_blocks.append(vision_text.strip())
            except Exception as e:
                logger.debug("Vision DOCX extraction failed on image %d: %s", idx + 1, e)
                continue
    except Exception as e:
        logger.warning("Vision DOCX extraction failed for %s: %s", basename, e)
        return []

    logger.info(
        "Vision DOCX extraction done - file=%s, images_augmented=%d",
        basename, len(vision_blocks),
    )
    return vision_blocks


# ── Table extraction helpers ─────────────────────────────────────────────


def _table_to_markdown(headers: list[str], rows: list[list[str]]) -> str:
    """Convert table headers + rows into clean markdown format."""
    # Pad cells to align columns
    all_rows = [headers] + rows
    col_widths = [max(len(str(cell)) for cell in col) for col in zip(*all_rows)] if all_rows else []

    def format_row(cells):
        parts = []
        for cell, w in zip(cells, col_widths):
            parts.append(str(cell).ljust(w))
        return "| " + " | ".join(parts) + " |"

    lines = [format_row(headers)]
    lines.append("| " + " | ".join("-" * w for w in col_widths) + " |")
    for row in rows:
        # Pad row if it has fewer columns than headers
        padded = row + [""] * (len(headers) - len(row))
        lines.append(format_row(padded))
    return "\n".join(lines)


def _extract_tables_from_pdf(file_path: str) -> dict[int, list[str]]:
    """
    Extract tables from each page of a PDF using PyMuPDF's find_tables().

    Returns: {page_index: [table_markdown_str, ...]}
    """
    import fitz

    tables_by_page = {}
    doc = fitz.open(file_path)

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        tab_finder = page.find_tables()

        if not tab_finder.tables:
            continue

        page_tables = []
        for table in tab_finder.tables:
            try:
                data = table.extract()
                if not data or len(data) < 2:
                    continue

                # First row as headers
                headers = [str(cell).strip() if cell else "" for cell in data[0]]
                rows = []
                for row in data[1:]:
                    rows.append([str(cell).strip() if cell else "" for cell in row])

                # Skip tiny tables (1 row, 1 col — likely noise)
                if len(rows) < 1 or len(headers) < 2:
                    continue

                md = _table_to_markdown(headers, rows)
                page_tables.append(md)
            except Exception as e:
                logger.debug("Table extraction failed on page %d: %s", page_idx, e)
                continue

        if page_tables:
            tables_by_page[page_idx] = page_tables
            logger.debug("Page %d — %d table(s) extracted", page_idx, len(page_tables))

    doc.close()
    logger.info(
        "PDF table extraction done — file=%s, pages_with_tables=%d",
        os.path.basename(file_path), len(tables_by_page),
    )
    return tables_by_page


def _extract_tables_from_docx(file_path: str) -> list[str]:
    """
    Extract all tables from a DOCX file.

    Returns: list of markdown-formatted table strings.
    """
    import docx

    doc = docx.Document(file_path)
    all_tables = []

    for table_idx, table in enumerate(doc.tables):
        try:
            if len(table.rows) < 2 or len(table.columns) < 2:
                continue

            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])

            md = _table_to_markdown(headers, rows)
            all_tables.append(md)
            logger.debug("DOCX table %d — %d rows x %d cols", table_idx, len(rows), len(headers))
        except Exception as e:
            logger.debug("DOCX table %d extraction failed: %s", table_idx, e)
            continue

    logger.info("DOCX table extraction — file=%s, tables=%d", os.path.basename(file_path), len(all_tables))
    return all_tables


def _build_table_document(
    table_md: str,
    source: str,
    page_number: int | None = None,
    table_index: int = 0,
    doc_id: str = "",
    filename: str = "",
) -> Document:
    """Create a Document for an extracted table with metadata."""
    meta = {
        "source": source,
        "doc_id": doc_id,
        "filename": filename,
        "is_table": True,
        "table_index": table_index,
    }
    if page_number is not None:
        meta["page_number"] = page_number

    # Prepend a header so the LLM knows this is a table
    content = f"[Table]\n{table_md}\n[/Table]"
    return Document(page_content=content, metadata=meta)


# ── Document loaders ─────────────────────────────────────────────────────


def _load_pdf(file_path: str) -> tuple[list[Document], dict[int, list[str]]]:
    """
    Load PDF pages as Documents, with tables extracted and formatted as
    markdown table chunks (marked with is_table=True metadata).

    Returns: (docs, tables_by_page) — tables_by_page is reused by caller
    to avoid double extraction.
    """
    from langchain_community.document_loaders import PyMuPDFLoader

    logger.debug("Loading PDF: %s", file_path)

    tables_by_page = _extract_tables_from_pdf(file_path)

    loader = PyMuPDFLoader(file_path)
    docs = loader.load()

    docs = apply_ocr_fallback(docs, file_path)

    vision_by_page = _extract_vision_text_from_pdf(file_path)

    for doc in docs:
        page_idx = doc.metadata.get("page", 0)
        if page_idx in tables_by_page:
            doc.metadata["has_tables"] = True
        if page_idx in vision_by_page:
            doc.page_content = _merge_extracted_and_vision_text(doc.page_content, vision_by_page[page_idx])
            doc.metadata["vision_augmented"] = True
            doc.metadata["vision_chars"] = len(vision_by_page[page_idx])

    logger.debug(
        "PDF loaded - pages=%d, pages_with_tables=%d, pages_with_vision=%d",
        len(docs), len(tables_by_page), len(vision_by_page),
    )
    return docs, tables_by_page


def _load_docx(file_path: str) -> list[Document]:
    """Load DOCX paragraphs as a Document, tables extracted separately."""
    import docx

    logger.debug("Loading DOCX: %s", file_path)
    doc = docx.Document(file_path)
    full_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    vision_blocks = _extract_vision_text_from_docx(file_path)
    merged_text = _merge_extracted_and_vision_text(full_text, "\n\n".join(vision_blocks))
    metadata = {
        "source": os.path.basename(file_path),
        "vision_augmented": bool(vision_blocks),
        "vision_blocks": len(vision_blocks),
    }
    logger.debug(
        "DOCX loaded - paragraphs=%d, chars=%d, vision_blocks=%d",
        len(doc.paragraphs), len(merged_text), len(vision_blocks),
    )
    return [Document(page_content=merged_text, metadata=metadata)]


def ingest_document(file_bytes: bytes, filename: str, doc_id: str | None = None) -> dict:
    import tempfile

    if doc_id is None:
        doc_id = str(uuid.uuid4())[:12]
    safe_name = safe_filename(filename)
    ext = os.path.splitext(filename)[1].lower()
    logger.info("Starting ingestion — file=%s, ext=%s, doc_id=%s", filename, ext, doc_id)

    # ── Load document ──────────────────────────────────────────────────
    load_start = time.time()
    if ext in [".jpg", ".jpeg", ".png"]:
        image = Image.open(io.BytesIO(file_bytes))

        # Always use vision model for images to extract chart data
        # Vision model provides much richer information than OCR alone
        logger.info("Using vision model for image: %s", filename)
        chart_description = _describe_chart_with_vision(image, filename)

        if chart_description and len(chart_description.strip()) > 50:
            docs = [Document(
                page_content=f"[Image: {filename}]\n{chart_description}",
                metadata={"source": filename, "is_image": True, "image_type": "chart"},
            )]
        else:
            # Fallback to OCR if vision model fails
            logger.warning("Vision model failed, falling back to OCR for: %s", filename)
            ocr_configs = [
                "--psm 3",   # Fully automatic page segmentation
                "--psm 6",   # Assume uniform block of text
                "--psm 11",  # Sparse text without order
                "--psm 12",  # Sparse text with order
            ]

            best_text = ""
            for config in ocr_configs:
                try:
                    text = pytesseract.image_to_string(image, config=config)
                    if len(text.strip()) > len(best_text.strip()):
                        best_text = text
                except Exception:
                    continue

            docs = [Document(
                page_content=f"[Image: {filename}]\n{best_text}",
                metadata={"source": filename, "is_image": True, "image_type": "text_document"},
            )]
        table_docs = []
    else:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            if ext == ".pdf":
                docs, tables_by_page = _load_pdf(tmp_path)
                table_docs = _build_table_docs_from_pages(tables_by_page, doc_id, filename)
            elif ext == ".docx":
                docs = _load_docx(tmp_path)
                table_docs = _extract_tables_as_docs_docx(tmp_path, doc_id, filename)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        finally:
            os.unlink(tmp_path)
    load_elapsed = time.time() - load_start
    logger.info(
        "Document loaded — pages=%d, table_chunks=%d, elapsed=%.2fs",
        len(docs), len(table_docs), load_elapsed,
    )

    # ── Split text into chunks (tables are kept separate) ──────────────
    chunk_size = int(os.getenv("CHUNK_SIZE", "512"))
    chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "50"))
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " "],
    )
    split_start = time.time()
    chunks = splitter.split_documents(docs)
    split_elapsed = time.time() - split_start
    logger.info(
        "Split into chunks — count=%d, chunk_size=%d, overlap=%d, elapsed=%.2fs",
        len(chunks), chunk_size, chunk_overlap, split_elapsed,
    )

    # ── Build embeddings + store ───────────────────────────────────────
    embeddings = get_embeddings()

    documents = []
    metadatas = []
    ids = []

    # 1) Text chunks
    total_chunks = len(chunks) + len(table_docs)
    for i, chunk in enumerate(chunks):
        documents.append(chunk.page_content)
        meta = chunk.metadata.copy()

        meta["doc_id"] = doc_id
        meta["filename"] = filename
        meta["chunk_index"] = i
        meta["total_chunks"] = total_chunks
        meta["is_table"] = False

        if "page" in meta and isinstance(meta["page"], int):
            meta["page_number"] = meta["page"] + 1

        first_line = chunk.page_content.split("\n")[0].strip()
        if len(first_line) < 100 and (
            first_line.isupper()
            or first_line.endswith(":")
            or first_line.startswith("#")
            or first_line.startswith("Section")
            or first_line.startswith("Article")
        ):
            meta["section"] = first_line
        elif "section" not in meta:
            meta["section"] = ""

        meta["content_preview"] = chunk.page_content[:150].replace("\n", " ").strip()
        metadatas.append(meta)
        ids.append(f"{doc_id}::chunk_{i}")

    # 2) Table chunks (atomic — not split further)
    for j, tdoc in enumerate(table_docs):
        documents.append(tdoc.page_content)
        meta = tdoc.metadata.copy()
        meta["doc_id"] = doc_id
        meta["filename"] = filename
        meta["chunk_index"] = len(chunks) + j
        meta["total_chunks"] = total_chunks
        meta["is_table"] = True
        meta["content_preview"] = tdoc.page_content[:150].replace("\n", " ").strip()
        metadatas.append(meta)
        ids.append(f"{doc_id}::table_{j}")

    store_start = time.time()
    vectorstore.add_documents(documents=documents, metadatas=metadatas, ids=ids)
    refresh_bm25()
    store_elapsed = time.time() - store_start
    logger.info(
        "Stored in Qdrant — text_chunks=%d, table_chunks=%d, elapsed=%.2fs, total=%d",
        len(chunks), len(table_docs), store_elapsed, vectorstore.get_doc_count(),
    )

    return {
        "doc_id": doc_id,
        "filename": safe_name,
        "num_chunks": len(chunks) + len(table_docs),
        "num_tables": len(table_docs),
        "status": "success",
    }


def _extract_tables_as_docs_pdf(file_path: str, doc_id: str, filename: str) -> list[Document]:
    """Extract PDF tables and return them as individual Documents."""
    tables_by_page = _extract_tables_from_pdf(file_path)
    return _build_table_docs_from_pages(tables_by_page, doc_id, filename)


def _build_table_docs_from_pages(
    tables_by_page: dict[int, list[str]],
    doc_id: str,
    filename: str,
) -> list[Document]:
    """Build Document objects from already-extracted table markdown strings."""
    result = []
    for page_idx, page_tables in tables_by_page.items():
        for table_idx, table_md in enumerate(page_tables):
            result.append(_build_table_document(
                table_md=table_md,
                source=filename,
                page_number=page_idx + 1,
                table_index=table_idx,
                doc_id=doc_id,
                filename=filename,
            ))
    return result


def _extract_tables_as_docs_docx(file_path: str, doc_id: str, filename: str) -> list[Document]:
    """Extract DOCX tables and return them as individual Documents."""
    table_mds = _extract_tables_from_docx(file_path)
    result = []
    for idx, table_md in enumerate(table_mds):
        result.append(_build_table_document(
            table_md=table_md,
            source=file_path,
            table_index=idx,
            doc_id=doc_id,
            filename=filename,
        ))
    return result